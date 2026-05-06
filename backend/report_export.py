"""
Report export — PDF (reportlab + DejaVu) and DOCX (python-docx).

Renders a full CAP-style anatomopathology report from the dict shape we send
to the frontend. Covers identity, clinical context, specimens, microscopy
(Histo-A vs Histo-B dual read), cross-slide synthesis, debate, primary
diagnosis with staging, IHC panel, recommendations, literature, audit and
sign-off.
"""

from __future__ import annotations

import io
import json
import re
from datetime import datetime, timezone, timedelta
from typing import Any

from reportlab.lib.colors import HexColor, white
from reportlab.lib.enums import TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle, HRFlowable,
    KeepTogether, PageBreak,
)


_FONTS_REGISTERED = False
_DEJAVU = "/usr/share/fonts/truetype/dejavu"


def _register_fonts() -> None:
    global _FONTS_REGISTERED
    if _FONTS_REGISTERED:
        return
    pdfmetrics.registerFont(TTFont("DejaVu",        f"{_DEJAVU}/DejaVuSans.ttf"))
    pdfmetrics.registerFont(TTFont("DejaVu-Bold",   f"{_DEJAVU}/DejaVuSans-Bold.ttf"))
    pdfmetrics.registerFont(TTFont("DejaVu-Italic", f"{_DEJAVU}/DejaVuSans-Oblique.ttf"))
    pdfmetrics.registerFont(TTFont("DejaVu-Mono",   f"{_DEJAVU}/DejaVuSansMono.ttf"))
    _FONTS_REGISTERED = True


# ──────────────────────────────────────────────────────────────────────────────
INK         = HexColor("#1c1a16")
INK_SOFT    = HexColor("#4a4538")
MUTED       = HexColor("#807866")
RULE        = HexColor("#c8c1b1")
RULE_STRONG = HexColor("#5e574b")
ACCENT      = HexColor("#6b1d1d")
PAPER_2     = HexColor("#ebe6db")
PAPER_3     = HexColor("#f5f1e8")
DANGER      = HexColor("#7a2222")
OK          = HexColor("#2f5d3a")


def _styles() -> dict:
    base = getSampleStyleSheet()["Normal"]
    return {
        "title":    ParagraphStyle("title",  parent=base, fontName="DejaVu-Bold",
                                   fontSize=18, leading=22, textColor=INK, spaceAfter=2),
        "h1":       ParagraphStyle("h1",     parent=base, fontName="DejaVu-Bold",
                                   fontSize=13, leading=16, textColor=INK, spaceBefore=14, spaceAfter=4),
        "h2":       ParagraphStyle("h2",     parent=base, fontName="DejaVu-Bold",
                                   fontSize=11, leading=14, textColor=INK, spaceBefore=8, spaceAfter=2),
        "smcaps":   ParagraphStyle("smcaps", parent=base, fontName="DejaVu-Mono",
                                   fontSize=8, leading=10, textColor=MUTED, spaceAfter=2),
        "smcaps_d": ParagraphStyle("smcaps_d", parent=base, fontName="DejaVu-Mono",
                                   fontSize=8, leading=10, textColor=INK, spaceAfter=2),
        "body":     ParagraphStyle("body",   parent=base, fontName="DejaVu",
                                   fontSize=9.5, leading=13, textColor=INK, alignment=TA_LEFT, spaceAfter=3),
        "body_sm":  ParagraphStyle("body_sm", parent=base, fontName="DejaVu",
                                   fontSize=8.5, leading=11, textColor=INK_SOFT, spaceAfter=2),
        "italic":   ParagraphStyle("italic", parent=base, fontName="DejaVu-Italic",
                                   fontSize=9.5, leading=13, textColor=INK_SOFT, spaceAfter=3),
        "mono":     ParagraphStyle("mono",   parent=base, fontName="DejaVu-Mono",
                                   fontSize=8.5, leading=11, textColor=INK_SOFT),
        "mono_sm":  ParagraphStyle("mono_sm", parent=base, fontName="DejaVu-Mono",
                                   fontSize=7.5, leading=10, textColor=MUTED),
        "accent":   ParagraphStyle("accent", parent=base, fontName="DejaVu-Bold",
                                   fontSize=10, leading=13, textColor=ACCENT),
        "diag":     ParagraphStyle("diag",   parent=base, fontName="DejaVu-Bold",
                                   fontSize=14, leading=18, textColor=ACCENT, spaceAfter=6),
        "right":    ParagraphStyle("right",  parent=base, fontName="DejaVu-Mono",
                                   fontSize=8, leading=10, textColor=MUTED, alignment=TA_RIGHT),
    }


# ── helpers ──────────────────────────────────────────────────────────────────

def _safe(v: Any, default: str = "—") -> str:
    if v is None:
        return default
    s = str(v).strip()
    return s if s else default


def _strip_json_fences(text: str) -> str:
    """LLM outputs often wrap JSON in ```json fences. Strip them."""
    text = text.strip()
    text = re.sub(r"^```(?:json)?\s*", "", text)
    text = re.sub(r"\s*```$", "", text)
    return text.strip()


def _parse_findings(raw: str) -> dict:
    """Parse a Histo-* findings string. Returns {} if not JSON."""
    if not raw:
        return {}
    try:
        return json.loads(_strip_json_fences(raw))
    except Exception:
        return {}


def _flatten(v: Any, max_len: int = 600) -> str:
    if v is None:
        return "—"
    if isinstance(v, str):
        return v[:max_len].strip() or "—"
    if isinstance(v, (list, tuple)):
        return " · ".join(_flatten(x, max_len) for x in v if x)[:max_len] or "—"
    if isinstance(v, dict):
        try:
            return json.dumps(v, ensure_ascii=False)[:max_len]
        except Exception:
            return str(v)[:max_len]
    return str(v)[:max_len]


def _fmt_paris_now() -> str:
    paris = timezone(timedelta(hours=2))
    return datetime.now(paris).strftime("%d/%m/%Y · %H:%M Paris")


# ── PDF building blocks ──────────────────────────────────────────────────────

def _kv_table(rows: list[tuple[str, str]], styles: dict, col_w: tuple[int, int] = (55, 110)) -> Table:
    data = [[Paragraph(k, styles["smcaps"]), Paragraph(_safe(v), styles["body"])] for k, v in rows]
    t = Table(data, colWidths=[col_w[0] * mm, col_w[1] * mm])
    t.setStyle(TableStyle([
        ("BACKGROUND",   (0, 0), (-1, -1), PAPER_2),
        ("BOX",          (0, 0), (-1, -1), 0.5, RULE),
        ("INNERGRID",    (0, 0), (-1, -1), 0.25, RULE),
        ("VALIGN",       (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING",  (0, 0), (-1, -1), 7),
        ("RIGHTPADDING", (0, 0), (-1, -1), 7),
        ("TOPPADDING",   (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING",(0, 0), (-1, -1), 5),
    ]))
    return t


def _two_col_table(left_cells: list, right_cells: list, styles: dict, headers: tuple[str, str]) -> Table:
    """Side-by-side two-column comparison (Histo-A | Histo-B)."""
    head = [
        Paragraph(f"<b>{headers[0]}</b>", styles["body"]),
        Paragraph(f"<b>{headers[1]}</b>", styles["body"]),
    ]
    rows = [head]
    for a, b in zip(left_cells, right_cells):
        rows.append([a, b])
    t = Table(rows, colWidths=[82 * mm, 82 * mm])
    t.setStyle(TableStyle([
        ("BACKGROUND",   (0, 0), (-1, 0), RULE_STRONG),
        ("TEXTCOLOR",    (0, 0), (-1, 0), white),
        ("BOX",          (0, 0), (-1, -1), 0.5, RULE_STRONG),
        ("INNERGRID",    (0, 0), (-1, -1), 0.25, RULE),
        ("VALIGN",       (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING",  (0, 0), (-1, -1), 7),
        ("RIGHTPADDING", (0, 0), (-1, -1), 7),
        ("TOPPADDING",   (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING",(0, 0), (-1, -1), 5),
        ("BACKGROUND",   (0, 1), (-1, -1), PAPER_3),
    ]))
    return t


def _ihc_table(biomarkers: list[str], styles: dict) -> Table:
    """IHC panel — placeholder rows for each requested marker."""
    head = [
        Paragraph("<b>Marqueur</b>",       styles["body"]),
        Paragraph("<b>% cellules</b>",     styles["body"]),
        Paragraph("<b>Intensité</b>",      styles["body"]),
        Paragraph("<b>Interprétation</b>", styles["body"]),
    ]
    rows = [head]
    for m in biomarkers:
        rows.append([
            Paragraph(_safe(m), styles["body"]),
            Paragraph("À doser", styles["italic"]),
            Paragraph("À doser", styles["italic"]),
            Paragraph("À compléter par la plateforme IHC", styles["italic"]),
        ])
    t = Table(rows, colWidths=[40 * mm, 28 * mm, 28 * mm, 70 * mm])
    t.setStyle(TableStyle([
        ("BACKGROUND",   (0, 0), (-1, 0), RULE_STRONG),
        ("TEXTCOLOR",    (0, 0), (-1, 0), white),
        ("BOX",          (0, 0), (-1, -1), 0.5, RULE_STRONG),
        ("INNERGRID",    (0, 0), (-1, -1), 0.25, RULE),
        ("VALIGN",       (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING",  (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING",   (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING",(0, 0), (-1, -1), 5),
        ("BACKGROUND",   (0, 1), (-1, -1), PAPER_3),
    ]))
    return t


def _papers_block(title: str, papers: list[dict], accent: bool, styles: dict) -> list:
    if not papers:
        return [Paragraph(f"{title} — aucune.", styles["italic"]), Spacer(1, 4)]
    flow = [Paragraph(title, styles["smcaps_d"]), Spacer(1, 2)]
    for p in papers:
        ref = p.get("pmid") or ""
        src = "TCGA" if p.get("source") == "tcga_case" else "PMID"
        head = f"<b>{src} {ref}</b> · τ {float(p.get('score', 0)):.2f} — {_safe(p.get('title'), '')}"
        flow.append(Paragraph(head, styles["accent"] if accent else styles["body"]))
        meta = " · ".join(filter(None, [p.get("journal"), str(p.get("year") or ""), p.get("authors")]))
        if meta:
            flow.append(Paragraph(meta, styles["mono"]))
        if p.get("relevance"):
            flow.append(Paragraph(p["relevance"], styles["italic"]))
        if p.get("snippet"):
            flow.append(Paragraph(p["snippet"][:380], styles["body_sm"]))
        if p.get("url"):
            flow.append(Paragraph(p["url"], styles["mono_sm"]))
        flow.append(Spacer(1, 5))
    return flow


def _per_slide_summary(findings: dict) -> str:
    """Distill a Histo-* JSON output into a 1-3 line summary for the dual-read column."""
    if not findings:
        return "Lecture indisponible."
    parts = []
    if (dx := findings.get("primary_diagnosis") or findings.get("diagnosis")):
        parts.append(f"<b>{_flatten(dx, 200)}</b>")
    if (grade := findings.get("grade")):
        parts.append(f"Grade {_flatten(grade, 80)}")
    if (mit := findings.get("mitotic_index")):
        parts.append(f"Mitoses {_flatten(mit, 80)}")
    if (margin := findings.get("margin_status") or findings.get("margins")):
        parts.append(f"Marges {_flatten(margin, 80)}")
    if (lvi := findings.get("lvi") or findings.get("lymphovascular_invasion")):
        parts.append(f"LVI {_flatten(lvi, 60)}")
    if (pni := findings.get("pni") or findings.get("perineural_invasion")):
        parts.append(f"PNI {_flatten(pni, 60)}")
    if (notes := findings.get("notes") or findings.get("microscopic_findings") or findings.get("findings")):
        parts.append(_flatten(notes, 320))
    if not parts:
        return _flatten(findings, 380)
    return "<br/>".join(parts)


# ── main PDF renderer ────────────────────────────────────────────────────────

def render_pdf(report: dict[str, Any], patient_label: str = "") -> bytes:
    """Build a CAP-style PDF report. Returns the PDF as bytes."""
    _register_fonts()
    s = _styles()

    cap = report.get("cap_report") or {}

    diagnosis  = (cap.get("primary_diagnosis") or report.get("diagnosis") or "Diagnostic indéterminé")
    icd        = _safe(cap.get("icd_o_code"))
    pt         = _safe(cap.get("pt_stage"))
    pn         = _safe(cap.get("pn_stage"))
    margin     = _safe(cap.get("margin_status"))
    pni        = _safe(cap.get("perineural_invasion"))
    lvi        = _safe(cap.get("lymphovascular_invasion"))
    grade      = _safe(cap.get("grade"))
    mitoses    = _safe(cap.get("mitotic_index"))
    conf       = float(report.get("confidence") or 0)

    biomarkers = report.get("biomarkers") or cap.get("biomarkers") or []
    findings_k = cap.get("key_findings") or []
    recos      = cap.get("recommendations") or report.get("recommendations") or []
    debate_sum = report.get("debate_summary") or cap.get("debate_summary")
    debate     = report.get("debate_rounds") or cap.get("debate_rounds") or []

    histo_a = report.get("histo_a_results") or []
    histo_b = report.get("histo_b_results") or []
    cross   = report.get("cross_slide") or {}
    triage  = report.get("triage_results") or []

    clinical = report.get("clinical_data") or {}
    clinical_ctx = clinical.get("context") or ""
    clinical_age = clinical.get("age")

    slide_paths = report.get("slide_paths") or []
    case_id     = report.get("case_id") or "—"
    patient_id  = report.get("patient_id") or "—"

    lit         = report.get("literature") or {}
    lit_summary = lit.get("key_findings") or ""
    used        = lit.get("used_papers") or []
    suggested   = lit.get("suggested_papers") or []
    similar     = lit.get("similar_cases") or 0

    warnings = report.get("warnings") or []

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        topMargin=18 * mm, bottomMargin=16 * mm,
        leftMargin=18 * mm, rightMargin=18 * mm,
        title="PathMind — Rapport CAP",
        author="PathMind v0.2",
    )

    story: list = []

    # ── Header
    header = Table(
        [[
            Paragraph("<b>PathMind</b>", s["title"]),
            Paragraph(
                f"Rapport anatomopathologique<br/>"
                f"{_fmt_paris_now()}<br/>"
                f"Dossier {case_id}",
                s["right"],
            ),
        ]],
        colWidths=[100 * mm, 74 * mm],
    )
    header.setStyle(TableStyle([
        ("VALIGN",      (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING",(0, 0), (-1, -1), 0),
    ]))
    story.append(header)
    story.append(HRFlowable(width="100%", thickness=1.2, color=RULE_STRONG, spaceBefore=4, spaceAfter=10))

    # ── 1. Identité patient
    story.append(Paragraph("1. Identité patient", s["h1"]))
    story.append(_kv_table([
        ("Patient",       patient_label or "—"),
        ("Identifiant",   patient_id),
        ("Âge",           f"{clinical_age} ans" if clinical_age else "—"),
        ("Numéro de cas", case_id),
        ("Confiance multi-agents τ", f"{conf:.2f}"),
    ], s))

    # ── 2. Renseignements cliniques
    story.append(Paragraph("2. Renseignements cliniques", s["h1"]))
    story.append(Paragraph(_safe(clinical_ctx, "Aucun contexte clinique transmis."), s["body"]))

    # ── 3. Spécimens reçus
    story.append(Paragraph("3. Spécimens reçus", s["h1"]))
    if slide_paths:
        spec_rows = []
        for i, p in enumerate(slide_paths):
            triage_row = next((t for t in triage if t.get("slide_index") == i), None)
            tile_count = triage_row.get("tile_count") if triage_row else "—"
            mpp = triage_row.get("mpp_x") if triage_row else None
            obj = triage_row.get("objective_power") if triage_row else None
            mpp_s = f"{mpp:.3f} µm/px" if isinstance(mpp, (int, float)) else "—"
            obj_s = f"×{obj:.0f}" if isinstance(obj, (int, float)) else "—"
            spec_rows.append([
                Paragraph(f"L{i+1}", s["mono"]),
                Paragraph(p.split("/")[-1], s["body_sm"]),
                Paragraph(str(tile_count), s["mono"]),
                Paragraph(mpp_s, s["mono"]),
                Paragraph(obj_s, s["mono"]),
            ])
        head = [Paragraph(f"<b>{h}</b>", s["body"]) for h in ("Lame", "Fichier", "Tuiles", "Résolution", "Obj.")]
        rows = [head] + spec_rows
        t = Table(rows, colWidths=[14*mm, 90*mm, 18*mm, 26*mm, 18*mm])
        t.setStyle(TableStyle([
            ("BACKGROUND",   (0, 0), (-1, 0), RULE_STRONG),
            ("TEXTCOLOR",    (0, 0), (-1, 0), white),
            ("BOX",          (0, 0), (-1, -1), 0.5, RULE_STRONG),
            ("INNERGRID",    (0, 0), (-1, -1), 0.25, RULE),
            ("VALIGN",       (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING",  (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING",   (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING",(0, 0), (-1, -1), 4),
            ("BACKGROUND",   (0, 1), (-1, -1), PAPER_3),
        ]))
        story.append(t)
    else:
        story.append(Paragraph("Aucune lame transmise.", s["italic"]))

    # ── 4. Microscopie — dual read Histo-A vs Histo-B
    story.append(Paragraph("4. Microscopie — double lecture indépendante", s["h1"]))
    story.append(Paragraph(
        "Chaque lame a été relue indépendamment par deux modèles "
        "(Histo-A : Qwen2.5-VL-72B ; Histo-B : Meditron-70B). "
        "Les divergences sont arbitrées plus bas par le Chief Pathologist agent.",
        s["italic"],
    ))
    if histo_a or histo_b:
        n = max(len(histo_a), len(histo_b))
        left = []
        right = []
        for i in range(n):
            a = histo_a[i] if i < len(histo_a) else {}
            b = histo_b[i] if i < len(histo_b) else {}
            af = _parse_findings((a or {}).get("findings", ""))
            bf = _parse_findings((b or {}).get("findings", ""))
            left.append(Paragraph(
                f"<b>Lame L{i+1}</b><br/>" + _per_slide_summary(af),
                s["body_sm"],
            ))
            right.append(Paragraph(
                f"<b>Lame L{i+1}</b><br/>" + _per_slide_summary(bf),
                s["body_sm"],
            ))
        story.append(Spacer(1, 4))
        story.append(_two_col_table(left, right, s,
                                    headers=("Histo-A · Qwen2.5-VL-72B", "Histo-B · Meditron-70B")))
    else:
        story.append(Paragraph("Lectures Histo-A / Histo-B indisponibles.", s["italic"]))

    # ── 5. Synthèse cross-slide
    if cross:
        story.append(Paragraph("5. Synthèse cross-slide", s["h1"]))
        story.append(_kv_table([
            ("Pattern dominant",     _flatten(cross.get("dominant_pattern"), 240)),
            ("Lames affectées",      _flatten(cross.get("affected_slides"), 120)),
            ("Synthèse Histo-A",     _flatten(cross.get("synthesis_a"), 600)),
            ("Synthèse Histo-B",     _flatten(cross.get("synthesis_b"), 600)),
        ], s))
        if cross.get("disagreements"):
            story.append(Paragraph("Désaccords identifiés", s["h2"]))
            for d in cross["disagreements"]:
                story.append(Paragraph(f"› {_flatten(d, 600)}", s["body"]))

    # ── 6. Débat Chief — round-by-round arbitration
    if debate:
        story.append(Paragraph("6. Débat d'arbitrage (Chief Pathologist)", s["h1"]))
        for i, d in enumerate(debate, 1):
            who = "Histo-A" if d.get("agent_id") == "histo_a" else "Histo-B"
            color = ACCENT if who == "Histo-A" else RULE_STRONG
            head = f'<font color="{color.hexval()}"><b>R{i} · {who}</b></font>'
            arg = _flatten(d.get("argument"), 900)
            ceded = " (a concédé)" if d.get("conceded") else ""
            story.append(Paragraph(f"{head}{ceded}", s["body"]))
            story.append(Paragraph(arg, s["body_sm"]))
            story.append(Spacer(1, 3))
    if debate_sum:
        story.append(Paragraph("Synthèse de l'arbitrage", s["h2"]))
        story.append(Paragraph(_flatten(debate_sum, 1800), s["body"]))

    # ── 7. Diagnostic primaire
    story.append(PageBreak())
    story.append(Paragraph("7. Diagnostic primaire", s["h1"]))
    story.append(Paragraph(_flatten(diagnosis, 400), s["diag"]))
    story.append(_kv_table([
        ("ICD-O-3",                     icd),
        ("Stade pT",                    pt),
        ("Stade pN",                    pn),
        ("Marges",                      margin),
        ("Engainement périnerveux",     pni),
        ("Invasion lymphovasculaire",   lvi),
        ("Grade",                       grade),
        ("Index mitotique",             mitoses),
    ], s))

    # ── 8. Constatations clés
    if findings_k:
        story.append(Paragraph("8. Constatations clés", s["h1"]))
        for i, f in enumerate(findings_k, 1):
            story.append(Paragraph(f"{i}. {_flatten(f, 600)}", s["body"]))

    # ── 9. Panel IHC
    if biomarkers:
        story.append(Paragraph("9. Panel immunohistochimique recommandé", s["h1"]))
        story.append(Paragraph(
            "Marqueurs à doser sur la plateforme IHC. Les valeurs % cellules / intensité "
            "sont à compléter par le technicien après réalisation des lames.",
            s["italic"],
        ))
        story.append(Spacer(1, 4))
        story.append(_ihc_table(biomarkers, s))

    # ── 10. Recommandations cliniques
    if recos:
        story.append(Paragraph("10. Recommandations cliniques et conduite à tenir", s["h1"]))
        for r in recos:
            story.append(Paragraph(f"› {_flatten(r, 600)}", s["body"]))

    # ── 11. Littérature
    story.append(Paragraph("11. Revue de littérature", s["h1"]))
    story.append(Paragraph(
        f"Recherche sémantique RAG sur PubMed + cohorte TCGA. "
        f"{similar} cas similaires identifiés.",
        s["italic"],
    ))
    if lit_summary:
        story.append(Paragraph("Synthèse Literature Hunter", s["h2"]))
        story.append(Paragraph(_flatten(lit_summary, 1800), s["body"]))
    story.append(Spacer(1, 4))
    story.extend(_papers_block("Références citées par le Chief", used, accent=True, styles=s))
    story.extend(_papers_block("Suggestions complémentaires", suggested, accent=False, styles=s))

    # ── 12. Audit anti-hallucination
    if warnings:
        story.append(Paragraph("12. Audit anti-hallucination", s["h1"]))
        story.append(Paragraph(
            "Validateurs automatiques exécutés en post-pipeline (formatage TNM, "
            "cohérence ICD-O, présence des marqueurs IHC standards, etc.).",
            s["italic"],
        ))
        for w in warnings:
            sev = w.get("severity", "info")
            sev_label = {"danger": "À VÉRIFIER", "warn": "Attention", "info": "Info"}.get(sev, "Info")
            color = DANGER if sev == "danger" else (ACCENT if sev == "warn" else MUTED)
            head = f'<font color="{color.hexval()}"><b>[{sev_label}]</b></font> {_flatten(w.get("message"), 400)}'
            story.append(Paragraph(head, s["body"]))
            if w.get("evidence"):
                story.append(Paragraph(f"→ {_flatten(w['evidence'], 280)}", s["mono"]))
            story.append(Spacer(1, 3))

    # ── 13. Sign-off
    story.append(HRFlowable(width="100%", thickness=0.6, color=RULE_STRONG, spaceBefore=18, spaceAfter=8))
    story.append(Paragraph("13. Validation", s["h2"]))
    sign = Table(
        [[
            Paragraph("Anatomopathologiste référent", s["smcaps"]),
            Paragraph("Date / Signature", s["smcaps"]),
        ],
        [
            Paragraph("Dr. _____________________________", s["body"]),
            Paragraph("___ / ___ / ______", s["body"]),
        ]],
        colWidths=[90 * mm, 84 * mm],
    )
    sign.setStyle(TableStyle([
        ("VALIGN",       (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING",  (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
        ("TOPPADDING",   (0, 0), (-1, 0),  6),
        ("TOPPADDING",   (0, 1), (-1, 1),  18),
    ]))
    story.append(sign)

    story.append(Spacer(1, 12))
    story.append(Paragraph(
        "Rapport généré automatiquement par PathMind v0.2 — pipeline LangGraph multi-agents "
        "(Tile-Triage · Histo-A Qwen2.5-VL-72B · Histo-B Meditron-70B · Cross-Slide · "
        "Literature Hunter · Chief). Diagnostic à valider par un anatomopathologiste habilité. "
        "Document non opposable en l'état — usage hors situation clinique réelle.",
        s["italic"],
    ))

    doc.build(story)
    return buf.getvalue()


# ──────────────────────────────────────────────────────────────────────────────
#  DOCX
# ──────────────────────────────────────────────────────────────────────────────

def render_docx(report: dict[str, Any], patient_label: str = "") -> bytes:
    """Build a CAP-style DOCX report. Returns the DOCX as bytes."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    cap = report.get("cap_report") or {}
    diagnosis = cap.get("primary_diagnosis") or report.get("diagnosis") or "Diagnostic indéterminé"
    conf = float(report.get("confidence") or 0)

    clinical = report.get("clinical_data") or {}
    clinical_ctx = clinical.get("context") or ""
    clinical_age = clinical.get("age")

    histo_a = report.get("histo_a_results") or []
    histo_b = report.get("histo_b_results") or []
    cross   = report.get("cross_slide") or {}
    debate  = report.get("debate_rounds") or cap.get("debate_rounds") or []
    biomarkers = report.get("biomarkers") or cap.get("biomarkers") or []
    findings = cap.get("key_findings") or []
    recos    = cap.get("recommendations") or report.get("recommendations") or []
    warnings = report.get("warnings") or []
    lit      = report.get("literature") or {}

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    doc.add_heading("PathMind — Rapport CAP", level=0)
    doc.add_paragraph(
        f"{patient_label or '—'} · confiance multi-agents τ {conf:.2f} · {_fmt_paris_now()}"
    )

    # 1. Identité
    doc.add_heading("1. Identité patient", level=1)
    t = doc.add_table(rows=4, cols=2)
    t.style = "Light Grid Accent 1"
    rows = [
        ("Patient",       patient_label or "—"),
        ("Identifiant",   report.get("patient_id") or "—"),
        ("Âge",           f"{clinical_age} ans" if clinical_age else "—"),
        ("Numéro de cas", report.get("case_id") or "—"),
    ]
    for i, (k, v) in enumerate(rows):
        t.rows[i].cells[0].text = k
        t.rows[i].cells[1].text = v

    # 2. Clinique
    doc.add_heading("2. Renseignements cliniques", level=1)
    doc.add_paragraph(clinical_ctx or "Aucun contexte clinique transmis.")

    # 3. Spécimens
    doc.add_heading("3. Spécimens reçus", level=1)
    for i, p in enumerate(report.get("slide_paths") or []):
        doc.add_paragraph(f"L{i+1} · {p.split('/')[-1]}", style="List Bullet")

    # 4. Microscopie — dual read
    doc.add_heading("4. Microscopie — double lecture indépendante", level=1)
    n = max(len(histo_a), len(histo_b))
    if n:
        tab = doc.add_table(rows=n+1, cols=2)
        tab.style = "Light Grid Accent 1"
        tab.rows[0].cells[0].text = "Histo-A · Qwen2.5-VL-72B"
        tab.rows[0].cells[1].text = "Histo-B · Meditron-70B"
        for i in range(n):
            a = histo_a[i] if i < len(histo_a) else {}
            b = histo_b[i] if i < len(histo_b) else {}
            af = _parse_findings((a or {}).get("findings", ""))
            bf = _parse_findings((b or {}).get("findings", ""))
            tab.rows[i+1].cells[0].text = f"L{i+1} — " + _per_slide_summary(af).replace("<br/>", "\n").replace("<b>", "").replace("</b>", "")
            tab.rows[i+1].cells[1].text = f"L{i+1} — " + _per_slide_summary(bf).replace("<br/>", "\n").replace("<b>", "").replace("</b>", "")

    # 5. Cross-slide
    if cross:
        doc.add_heading("5. Synthèse cross-slide", level=1)
        for k, v in [
            ("Pattern dominant",  cross.get("dominant_pattern")),
            ("Lames affectées",   cross.get("affected_slides")),
            ("Synthèse Histo-A",  cross.get("synthesis_a")),
            ("Synthèse Histo-B",  cross.get("synthesis_b")),
        ]:
            doc.add_paragraph().add_run(f"{k} : ").bold = True
            doc.paragraphs[-1].add_run(_flatten(v, 600))
        if cross.get("disagreements"):
            doc.add_heading("Désaccords identifiés", level=2)
            for d in cross["disagreements"]:
                doc.add_paragraph(_flatten(d, 600), style="List Bullet")

    # 6. Débat
    if debate:
        doc.add_heading("6. Débat d'arbitrage", level=1)
        for i, d in enumerate(debate, 1):
            who = "Histo-A" if d.get("agent_id") == "histo_a" else "Histo-B"
            p = doc.add_paragraph()
            p.add_run(f"R{i} · {who} : ").bold = True
            p.add_run(_flatten(d.get("argument"), 900))
    if report.get("debate_summary"):
        doc.add_heading("Synthèse de l'arbitrage", level=2)
        doc.add_paragraph(report["debate_summary"])

    # 7. Diagnostic
    doc.add_heading("7. Diagnostic primaire", level=1)
    p = doc.add_paragraph()
    r = p.add_run(diagnosis)
    r.bold = True
    r.font.size = Pt(13)

    t2 = doc.add_table(rows=8, cols=2)
    t2.style = "Light Grid Accent 1"
    rows2 = [
        ("ICD-O-3",                  str(cap.get("icd_o_code") or "—")),
        ("Stade pT",                 str(cap.get("pt_stage") or "—")),
        ("Stade pN",                 str(cap.get("pn_stage") or "—")),
        ("Marges",                   str(cap.get("margin_status") or "—")),
        ("Engainement périnerveux",  str(cap.get("perineural_invasion") or "—")),
        ("Invasion lymphovasculaire",str(cap.get("lymphovascular_invasion") or "—")),
        ("Grade",                    str(cap.get("grade") or "—")),
        ("Index mitotique",          str(cap.get("mitotic_index") or "—")),
    ]
    for i, (k, v) in enumerate(rows2):
        t2.rows[i].cells[0].text = k
        t2.rows[i].cells[1].text = v

    # 8. Constatations
    if findings:
        doc.add_heading("8. Constatations clés", level=1)
        for f in findings:
            doc.add_paragraph(_flatten(f, 600), style="List Number")

    # 9. IHC
    if biomarkers:
        doc.add_heading("9. Panel immunohistochimique recommandé", level=1)
        ih = doc.add_table(rows=len(biomarkers)+1, cols=4)
        ih.style = "Light Grid Accent 1"
        for i, h in enumerate(["Marqueur", "% cellules", "Intensité", "Interprétation"]):
            ih.rows[0].cells[i].text = h
        for i, m in enumerate(biomarkers):
            ih.rows[i+1].cells[0].text = m
            ih.rows[i+1].cells[1].text = "À doser"
            ih.rows[i+1].cells[2].text = "À doser"
            ih.rows[i+1].cells[3].text = "À compléter par la plateforme IHC"

    # 10. Recommandations
    if recos:
        doc.add_heading("10. Recommandations cliniques et conduite à tenir", level=1)
        for r in recos:
            doc.add_paragraph(_flatten(r, 600), style="List Bullet")

    # 11. Littérature
    doc.add_heading("11. Revue de littérature", level=1)
    if lit.get("key_findings"):
        doc.add_paragraph(lit["key_findings"])
    for label, papers in [("Références citées par le Chief", lit.get("used_papers") or []),
                          ("Suggestions complémentaires",    lit.get("suggested_papers") or [])]:
        doc.add_heading(label, level=2)
        if not papers:
            doc.add_paragraph("Aucune.").italic = True
            continue
        for p in papers:
            src = "TCGA" if p.get("source") == "tcga_case" else "PMID"
            line = doc.add_paragraph(style="List Bullet")
            line.add_run(f"{src} {p.get('pmid','')} — ").bold = True
            line.add_run(p.get("title", ""))
            if p.get("journal") or p.get("year"):
                line.add_run(f"\n{p.get('journal','')} · {p.get('year','')}")
            if p.get("snippet"):
                line.add_run(f"\n{p['snippet'][:380]}").italic = True
            if p.get("url"):
                line.add_run(f"\n{p['url']}")

    # 12. Audit
    if warnings:
        doc.add_heading("12. Audit anti-hallucination", level=1)
        for w in warnings:
            sev = {"danger": "À VÉRIFIER", "warn": "Attention", "info": "Info"}.get(w.get("severity", "info"), "Info")
            p = doc.add_paragraph()
            p.add_run(f"[{sev}] ").bold = True
            p.add_run(_flatten(w.get("message"), 400))
            if w.get("evidence"):
                p.add_run(f"\n→ {_flatten(w['evidence'], 280)}").italic = True

    # 13. Sign-off
    doc.add_heading("13. Validation", level=1)
    doc.add_paragraph("Anatomopathologiste référent : Dr. _____________________________")
    doc.add_paragraph("Date / Signature : ___ / ___ / ______")

    p = doc.add_paragraph()
    r = p.add_run(
        "Rapport généré automatiquement par PathMind v0.2 — pipeline LangGraph multi-agents. "
        "Diagnostic à valider par un anatomopathologiste habilité. "
        "Document non opposable en l'état — usage hors situation clinique réelle."
    )
    r.italic = True
    r.font.color.rgb = RGBColor(0x80, 0x78, 0x66)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
